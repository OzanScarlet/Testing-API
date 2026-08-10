import os
import json
import time
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI, RateLimitError

load_dotenv()

MODEL_NAME = os.getenv("MODEL_NAME")
BASE_URL = os.getenv("BASE_URL")
DEFAULT_FOLDER = "data"
OUTPUT_DIR = "output"

# Semakin kecil chunk, semakin ringan request -> kecil kemungkinan kena limit.
CHUNK_CHARS = int(os.getenv("CHUNK_CHARS", "5000"))
MIN_CHUNK_CHARS = 2000

# Batasi jumlah pertanyaan per chunk agar output tidak boros token (lebih cepat).
MAX_QUESTIONS_PER_CHUNK = os.getenv("MAX_QUESTIONS_PER_CHUNK")

DELAY_BETWEEN_FILES = float(os.getenv("DELAY", "3.0"))
MAX_DELAY = 60.0
MAX_RETRIES = 6

current_delay = DELAY_BETWEEN_FILES


def get_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    return OpenAI(
        api_key=api_key,
        base_url=BASE_URL,
        default_headers={"User-Agent": "python-requests/2.32"},
    )


def read_documents(folder_path: str) -> list:
    """
    Membaca seluruh file .md di folder dan subfolder secara rekursif.
    Mengembalikan list berisi konten tiap file (bukan digabung).
    """
    documents = []
    path = Path(folder_path)
    md_files = list(path.rglob("*.md"))

    if not md_files:
        print(f"[WARN] Tidak ditemukan file .md di folder '{folder_path}'")
        return documents

    print(f"[INFO] Ditemukan {len(md_files)} file .md:")
    for file_path in md_files:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            documents.append(
                {
                    "source": file_path.name,
                    "path": str(file_path),
                    "content": content,
                }
            )
            print(f"[INFO] Membaca: {file_path}")
        except Exception as e:
            print(f"[ERROR] Gagal membaca file {file_path}: {e}")

    return documents


def split_document(text: str) -> list:
    """
    Memecah dokumen menjadi chunk kecil (per heading / per batas karakter)

    """
    lines = text.splitlines()
    chunks = []
    current = []
    current_len = 0

    for line in lines:
        is_heading = line.lstrip().startswith("#")
        force_split = current and current_len >= CHUNK_CHARS
        split_at_heading = is_heading and current and current_len >= MIN_CHUNK_CHARS

        if force_split or split_at_heading:
            chunks.append("\n".join(current).strip())
            current = []
            current_len = 0

        current.append(line)
        current_len += len(line) + 1

    if current:
        chunks.append("\n".join(current).strip())

    return [c for c in chunks if c.strip()]


def pace():
    time.sleep(current_delay)


def bump_delay():
    """Naikkan jeda secara adaptif."""
    global current_delay
    current_delay = min(current_delay * 2, MAX_DELAY)
    print(f"[RATE LIMIT] Delay dinaikkan menjadi {current_delay:.0f} detik.")


def wait_from_headers(e: RateLimitError, fallback: float) -> float:
    headers = getattr(getattr(e, "response", None), "headers", None) or {}

    def _to_secs(value):
        if not value:
            return None
        try:
            if str(value).lower().endswith("s"):
                return float(str(value)[:-1])
            return float(value)
        except (ValueError, TypeError):
            return None

    wait = _to_secs(headers.get("retry-after"))
    if wait is not None:
        return wait + 1

    if str(headers.get("x-ratelimit-remaining-tokens", "")) == "0":
        reset = _to_secs(headers.get("x-ratelimit-reset-tokens"))
        if reset is not None:
            return reset + 1

    if str(headers.get("x-ratelimit-remaining-requests", "")) == "0":
        reset = _to_secs(headers.get("x-ratelimit-reset-requests"))
        if reset is not None:
            return reset + 1

    return fallback


def generate_questions(chunk_text: str) -> list:
    """
    Mengirimkan satu chunk ke API untuk menghasilkan daftar PERTANYAAN saja
    . Output dibatasi agar cepat dan hemat token.
    """
    if not chunk_text.strip():
        return []

    if MAX_QUESTIONS_PER_CHUNK:
        limit_text = (
            f"Buat MAKSIMAL {MAX_QUESTIONS_PER_CHUNK} pertanyaan yang paling penting dan spesifik "
        )
    else:
        limit_text = (
            "Buat SEMUA pertanyaan yang paling penting dan spesifik (jangan batasi jumlah) "
        )

    system_prompt = (
        "Kamu adalah pakar pembuat soal dan evaluator sistem RAG. "
        "Tugasmu membaca potongan dokumen ketentuan/aturan lalu membuat daftar pertanyaan uji. "
        + limit_text
        + "yang bisa dijawab dari isi dokumen. "
        "Cukup daftar pertanyaan saja, TANPA jawaban. "
        "OUTPUT WAJIB BERFORMAT JSON DENGAN STRUKTUR KETAT TANPA KUNCI LAIN.\n"
        'Contoh output valid:\n'
        '{\n  "questions": ["Pertanyaan 1?", "Pertanyaan 2?"]\n}'
    )

    user_prompt = f"Berikut isi potongan dokumen:\n\n{chunk_text}"

    client = get_client()

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            pace()
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                response_format={"type": "json_object"},
                temperature=0.3,
            )

            raw_output = response.choices[0].message.content
            data = json.loads(raw_output)

            if "questions" in data and isinstance(data["questions"], list):
                return data["questions"]
            if isinstance(data, list):
                return data
            return list(data.values())[0]

        except RateLimitError as e:
            bump_delay()
            wait = wait_from_headers(e, fallback=2 ** attempt)
            print(f"[RATE LIMIT] Percobaan {attempt}/{MAX_RETRIES}. Tunggu {wait:.0f} detik...")
            time.sleep(wait)
        except Exception as e:
            print(f"[ERROR] Error saat memanggil API: {e}")
            return []

    print(f"[ERROR] Chunk gagal setelah {MAX_RETRIES} percobaan karena rate limit.")
    return []


def output_path_for(folder_path: str) -> str:
    path = Path(folder_path).resolve()
    default_root = Path(DEFAULT_FOLDER).resolve()
    if path == default_root:
        return os.path.join(OUTPUT_DIR, "questions.json")
    return os.path.join(OUTPUT_DIR, f"questions_{path.name.replace(' ', '_').lower()}.json")


def normalize_item(item: dict) -> dict:
    """
    Menormalkan item dari hasil lama (agar aman untuk resume).
    Format lama mungkin berisi "qas"/jawaban -> dibersihkan jadi pertanyaan saja.
    """
    item = dict(item)
    item.setdefault("source", Path(item.get("path", "")).name)

    questions = list(item.get("questions", []))
    if not questions and "qas" in item:
        questions = [qa["question"] for qa in item.get("qas", [])]

    item["questions"] = questions
    item.pop("qas", None)
    item.setdefault("done_chunks", [])
    return item


def load_existing_results(output_path: str) -> dict:
    """
    Resume: memuat hasil lama.
    - "completed": {path: item} -> file yang sukses penuh, dilewati.
    - "incomplete": {path: item} -> file yang gagal, hanya chunk yang belum
      selesai (dari 'done_chunks') yang diproses ulang, hemat token.
    """
    if not os.path.exists(output_path):
        return {"completed": {}, "incomplete": {}}

    try:
        with open(output_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        failed = set(data.get("failed_files", []))
        completed = {}
        incomplete = {}
        for item in data.get("files", []):
            item = normalize_item(item)
            if item["path"] in failed:
                incomplete[item["path"]] = item
            else:
                completed[item["path"]] = item
        return {"completed": completed, "incomplete": incomplete}
    except Exception as e:
        print(f"[WARN] Gagal membaca hasil lama ({output_path}): {e}. Mulai dari awal.")
        return {"completed": {}, "incomplete": {}}


def build_result(results: list, incomplete: dict) -> dict:
    all_questions = [q for item in results for q in item.get("questions", [])]
    files = results + list(incomplete.values())
    return {
        "total_files": len(files),
        "total_questions": len(all_questions),
        "questions": all_questions,
        "failed_files": sorted(incomplete.keys()),
        "files": files,
    }


def save_to_json(data, output_path: str) -> None:
    """Menyimpan hasil ke file JSON."""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"[OK] Hasil tersimpan di: {output_path}")


def process_folder(folder_path: str = DEFAULT_FOLDER, output_file: str = None) -> list:
    """
    Proses satu folder: dokumen -> chunk kecil -> pertanyaan per chunk.
    Hanya menghasilkan PERTANYAAN (tanpa jawaban).
    - Resume per-chunk: file sukses dilewati; file gagal hanya chunk yang
      belum selesai yang diproses ulang.
    - Progres disimpan setiap chunk (aman walau terputus).
    """
    print("=== QUESTIONS GENERATOR (PERTANYAAN SAJA, CHUNK KECIL) ===")
    print(f"[INFO] Folder: {folder_path}")

    if output_file is None:
        output_file = output_path_for(folder_path)

    documents = read_documents(folder_path)
    if not documents:
        print("[ERROR] Tidak ada dokumen yang bisa diproses.")
        return []

    state = load_existing_results(output_file)
    completed = state["completed"]
    incomplete = state["incomplete"]
    print(f"[INFO] File sudah selesai (dilewati): {len(completed)}")
    print(f"[INFO] File belum selesai (dilanjutkan): {len(incomplete)}")

    results = []

    for i, doc in enumerate(documents, start=1):
        path = doc["path"]
        source = doc["source"]

        if path in completed:
            print(f"[SKIP] {source} sudah selesai. Dilewati.")
            results.append(completed[path])
            continue

        chunks = split_document(doc["content"])

        item = incomplete.pop(path, None)
        if item is None:
            item = {"source": source, "path": path, "questions": [], "done_chunks": []}
        file_questions = list(item["questions"])
        done = set(item.get("done_chunks", []))
        pending = [idx for idx in range(len(chunks)) if idx not in done]

        print(f"\n[{i}/{len(documents)}] {source} -> {len(chunks)} chunk, sisa {len(pending)}")

        for idx in pending:
            c = idx + 1
            print(f"  - Chunk {c}/{len(chunks)} ({len(chunks[idx])} karakter)")
            questions = generate_questions(chunks[idx])
            if questions:
                file_questions.extend(questions)
                done.add(idx)
                item["questions"] = file_questions
                item["done_chunks"] = sorted(done)
                print(f"    -> {len(questions)} pertanyaan")
            else:
                print(f"    -> gagal, disimpan untuk run berikutnya")

            # Simpan progres setiap chunk (aman walau proses terputus)
            incomplete[path] = dict(item)
            save_to_json(build_result(results, incomplete), output_file)

        if len(done) == len(chunks) and file_questions:
            results.append(item)
            incomplete.pop(path, None)
            print(f"[OK] {source} -> {len(file_questions)} pertanyaan")
        else:
            incomplete[path] = dict(item)
            print(f"[FAIL] {source} belum lengkap, akan dilanjutkan di run berikutnya.")

        save_to_json(build_result(results, incomplete), output_file)

    result = build_result(results, incomplete)
    save_to_json(result, output_file)

    total = sum(len(r.get("questions", [])) for r in results)
    print(f"\n[OK] Selesai. {total} pertanyaan dari {len(results)} file selesai.")
    if incomplete:
        print(f"[INFO] {len(incomplete)} file masih berjalan. Jalankan ulang untuk melanjutkan.")
        for p in sorted(incomplete):
            print(f"       - {p}")
    print(f"[OK] Output: {output_file}")
    return result


def extract_text(file_path: str) -> str:
    """Ekstrak teks dari file: .md/.txt (langsung), .pdf (pypdf),
    .docx (python-docx), .doc (best-effort). Return str (bisa kosong)."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in (".md", ".txt"):
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="latin-1", errors="replace")

    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(pages)

    if suffix == ".docx":
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    if suffix == ".doc":
        # Best-effort: coba baca sebagai docx (kalau sebenarnya docx), lalu teks polos.
        try:
            import docx
            d = docx.Document(str(path))
            return "\n".join(p.text for p in d.paragraphs)
        except Exception:
            try:
                return path.read_text(encoding="latin-1", errors="replace")
            except Exception:
                return ""

    return ""


def iter_generate_questions_from_files(file_paths: list, output_file: str = None):
    """Generator STREAMING: proses banyak file (.md/.txt/.pdf/.docx/.doc)
    dan yield tiap hasil chunk begitu selesai digenerate.

    Yield tuple per chunk:
        (source, chunk_idx, chunk_total, list_questions, item, items)

    - source      : nama file
    - chunk_idx   : index chunk (0-based)
    - chunk_total : jumlah chunk di file itu
    - list_questions : daftar pertanyaan hasil generate chunk tsb
    - item        : item file yang sedang diisi (source/path/questions/done_chunks)
    - items       : list item yang sudah selesai sejauh ini

    Setiap chunk selesai, snapshot hasil (build_result + save_to_json) ditulis
    ke output_file (default output/questions_upload.json) supaya progres tidak
    hilang meski proses terputus.

    Caller bertanggung jawab mengevaluasi tiap pertanyaan dari list_questions
    sebelum meminta yield berikutnya."""
    print("=== QUESTIONS GENERATOR (STREAMING, MULTI FILE) ===")
    if output_file is None:
        output_file = os.path.join(OUTPUT_DIR, "questions_upload.json")

    items = []

    for file_path in file_paths:
        path = Path(file_path)
        print(f"[INFO] File: {path.name}")
        try:
            content = extract_text(str(path))
        except Exception as e:
            print(f"[ERROR] Gagal ekstrak teks {path.name}: {e}")
            continue

        if not content or not content.strip():
            print(f"[WARN] {path.name} tidak punya teks yang terbaca. Dilewati.")
            continue

        source = path.name
        chunks = split_document(content)
        item = {"source": source, "path": str(path), "questions": [], "done_chunks": []}

        for idx, chunk in enumerate(chunks, start=1):
            print(f"  - Chunk {idx}/{len(chunks)} ({len(chunk)} karakter)")
            questions = generate_questions(chunk)
            if questions:
                item["questions"].extend(questions)
                item["done_chunks"].append(idx - 1)
                print(f"    -> {len(questions)} pertanyaan")
                yield source, idx - 1, len(chunks), list(questions), dict(item), items
            else:
                print(f"    -> gagal, dilewati")

        if item["questions"]:
            items.append(item)
            print(f"[OK] {source} -> {len(item['questions'])} pertanyaan")
        else:
            print(f"[WARN] {source} tidak menghasilkan pertanyaan.")

        # Snapshot progres (semua file yang sudah selesai sejauh ini).
        if items:
            snapshot = build_result(items, {})
            save_to_json(snapshot, output_file)
            print(f"[OK] Snapshot sementara tersimpan: {output_file}")

    if not items:
        print("[ERROR] Tidak ada pertanyaan yang berhasil dihasilkan.")
    print("[DONE] Streaming selesai.")


def generate_questions_from_files(file_paths: list, output_file: str = None) -> dict:
    """
    Proses BANYAK file (.md/.txt/.pdf/.docx/.doc): ekstrak teks ->
    chunk kecil -> pertanyaan per chunk.
    Return dict berformat sama seperti hasil process_folder:
    {"total_files", "total_questions", "questions", "failed_files", "files"}.
    """
    if output_file is None:
        output_file = os.path.join(OUTPUT_DIR, "questions_upload.json")

    items = []
    for _source, _ci, _ct, _questions, item, items_so_far in iter_generate_questions_from_files(
        file_paths, output_file=output_file
    ):
        if item["questions"] and item not in items:
            items.append(item)

    if not items:
        return {}

    result = build_result(items, {})
    save_to_json(result, output_file)
    total = sum(len(i["questions"]) for i in items)
    print(f"[OK] {total} pertanyaan dari {len(items)} file selesai.")
    print(f"[OK] Output: {output_file}")
    return result


def generate_questions_from_file(file_path: str, output_file: str = None) -> dict:
    """Kompatibilitas: satu file -> pertanyaan (pakai multi-file)."""
    return generate_questions_from_files([file_path], output_file)


def main(folder_path: str = DEFAULT_FOLDER) -> list:
    """
    Jalankan seperti biasa: python modules/questions_generator.py
    atau per folder: python modules/questions_generator.py data/ksi
    Jalankan ulang bila masih ada file yang belum selesai (resume otomatis).
    """
    return process_folder(folder_path)


if __name__ == "__main__":
    main()
